"""
generate_report_v2.py — RC Credit Check Report Generator (v2)
Format aligned to actual RC team reports (8 changes vs v1).

Changes implemented:
1. Brief Profile → prose paragraphs (not bullets)
2. Financial Analysis → bold-labeled prose sub-sections
3. Strengths / Weaknesses → inline bold title in Normal paragraph
4. Latest Updates → bold "Month YYYY — " prefix in same run
5. View / Conditions → plain paragraphs, no separate "Conditions:" header
6. Scoring table → moved to very end of document
7. Header table → 4 fields only (Name, Incorporation, Business, Industry)
8. Requirement → printed at very top, before header table
"""

import json
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ──────────────────────────── helpers ────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), kwargs.get(edge, "nil"))
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get(f"{edge}Color", "auto"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def set_table_border(table, hex_color="000000", size="8"):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), hex_color)
        tblBorders.append(tag)
    tblPr.append(tblBorders)

def add_run(para, text, bold=False, italic=False, color=None, size=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    if size:
        run.font.size = Pt(size)
    return run

def heading(doc, text, level=1, color="1F3864"):
    """Section heading — dark navy, not using built-in Heading style."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11 if level == 1 else 10)
    run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def para(doc, text="", bold=False, italic=False, size=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
    return p

def divider(doc):
    """Thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    return p


# ──────────────────────────── banner ────────────────────────────

def add_banner(doc, company_name, date_str):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(4.5)
    table.columns[1].width = Inches(2.5)

    left = table.cell(0, 0)
    right = table.cell(0, 1)

    set_cell_bg(left, "1F3864")
    set_cell_bg(right, "1F3864")

    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr = lp.add_run(f"  CREDIT CHECK — {company_name.upper()}")
    lr.bold = True
    lr.font.size = Pt(13)
    lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run(f"Date: {date_str}  ")
    rr.font.size = Pt(9)
    rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()


# ──────────────────────────── header table (4 fields) ────────────────────────────

def add_header_table(doc, data):
    """
    Change #7: 4-field table only — Name, Incorporation, Business, Industry.
    """
    fields = [
        ("Company Name",       data.get("company_name", "")),
        ("Date of Incorporation", data.get("incorporation_date", "")),
        ("Nature of Business", data.get("nature_of_business", "")),
        ("Industry",           data.get("industry", "")),
    ]

    table = doc.add_table(rows=len(fields), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_border(table, "CCCCCC", "4")

    col_widths = [Inches(2.2), Inches(4.8)]
    for i, row in enumerate(table.rows):
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]
        set_cell_bg(row.cells[0], "EAF0FB")

        label_p = row.cells[0].paragraphs[0]
        label_p.paragraph_format.space_after = Pt(2)
        lr = label_p.add_run(fields[i][0])
        lr.bold = True
        lr.font.size = Pt(9)

        val_p = row.cells[1].paragraphs[0]
        val_p.paragraph_format.space_after = Pt(2)
        vr = val_p.add_run(fields[i][1])
        vr.font.size = Pt(9)

    doc.add_paragraph()


# ──────────────────────────── initial screening ────────────────────────────

def add_screening(doc, screening):
    heading(doc, "Initial Screening")
    table = doc.add_table(rows=1, cols=4)
    set_table_border(table, "AAAAAA", "4")
    headers = ["Parameter", "Threshold", "Actual", "Result"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_bg(cell, "1F3864")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row_data in screening:
        row = table.add_row()
        values = [
            row_data.get("parameter", ""),
            row_data.get("threshold", ""),
            row_data.get("actual", ""),
            row_data.get("result", ""),
        ]
        for i, val in enumerate(values):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            result = row_data.get("result", "")
            if i == 3:
                if "PASS" in result.upper():
                    r.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
                    r.bold = True
                elif "FAIL" in result.upper():
                    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                    r.bold = True

    doc.add_paragraph()


# ──────────────────────────── brief profile (prose) ────────────────────────────

def add_brief_profile(doc, paragraphs):
    """
    Change #1: prose paragraphs instead of bullets.
    `paragraphs` is a list of strings, each rendered as a Normal paragraph.
    """
    heading(doc, "Brief Profile")
    for text in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(text)
        r.font.size = Pt(10)
    doc.add_paragraph()


# ──────────────────────────── shareholding ────────────────────────────

def add_shareholding(doc, shareholders):
    heading(doc, "Shareholding Pattern")
    table = doc.add_table(rows=1, cols=2)
    set_table_border(table, "AAAAAA", "4")
    for i, h in enumerate(["Shareholder", "Stake"]):
        cell = table.cell(0, i)
        set_cell_bg(cell, "1F3864")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for sh in shareholders:
        row = table.add_row()
        for i, val in enumerate([sh.get("name", ""), sh.get("stake", "")]):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.size = Pt(9)

    doc.add_paragraph()


# ──────────────────────────── financial table ────────────────────────────

def add_financial_table(doc, financials):
    heading(doc, "Financial Snapshot (₹ Crores)")
    cols = financials.get("columns", [])
    rows_data = financials.get("rows", [])

    col_count = len(cols)
    table = doc.add_table(rows=1, cols=col_count)
    set_table_border(table, "AAAAAA", "4")

    header_row = table.rows[0]
    for i, col in enumerate(cols):
        cell = header_row.cells[i]
        set_cell_bg(cell, "1F3864")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(col)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for j, row_data in enumerate(rows_data):
        row = table.add_row()
        is_section = row_data.get("section_header", False)
        for i, val in enumerate(row_data.get("values", [])):
            cell = row.cells[i]
            if is_section:
                set_cell_bg(cell, "D9E1F2")
            elif j % 2 == 0:
                set_cell_bg(cell, "F5F7FC")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            r.bold = is_section or (i == 0)

    doc.add_paragraph()


# ──────────────────────────── financial analysis (prose) ────────────────────────────

def add_financial_analysis(doc, analysis_sections):
    """
    Change #2: bold-labeled prose sub-sections.
    `analysis_sections` is a list of {"label": "Revenue", "text": "..."}.
    Label is bolded inline at the start of the paragraph.
    """
    heading(doc, "Financial Analysis")
    for section in analysis_sections:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(0)
        label = section.get("label", "")
        text = section.get("text", "")
        if label:
            r_label = p.add_run(f"{label}: ")
            r_label.bold = True
            r_label.font.size = Pt(10)
        r_text = p.add_run(text)
        r_text.font.size = Pt(10)
    doc.add_paragraph()


# ──────────────────────────── strengths / weaknesses ────────────────────────────

def add_sw(doc, title, items):
    """
    Change #3: inline bold title within Normal paragraph.
    Each item: {"title": "Debt-free", "detail": "The company has..."}
    """
    heading(doc, title)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.space_before = Pt(0)
        t = item.get("title", "")
        d = item.get("detail", "")
        if t:
            r1 = p.add_run(f"{t} — ")
            r1.bold = True
            r1.font.size = Pt(10)
        r2 = p.add_run(d)
        r2.font.size = Pt(10)
    doc.add_paragraph()


# ──────────────────────────── latest updates ────────────────────────────

def add_latest_updates(doc, updates):
    """
    Change #4: bold "Month YYYY — " prefix in same paragraph run.
    Each update: {"period": "Feb 2026", "text": "..."}
    """
    heading(doc, "Latest Updates")
    for upd in updates:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.space_before = Pt(0)
        period = upd.get("period", "")
        text = upd.get("text", "")
        if period:
            r1 = p.add_run(f"{period} — ")
            r1.bold = True
            r1.font.size = Pt(10)
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
    doc.add_paragraph()


# ──────────────────────────── credit view & conditions ────────────────────────────

def add_view_and_conditions(doc, view_text, conditions, rating_label, decision):
    """
    Change #5: plain paragraphs, no "Conditions of Approval:" sub-heading.
    Change #6: scoring moved to end — this section just has view + conditions.
    """
    heading(doc, "Credit View")

    # Rating badge
    p_rating = doc.add_paragraph()
    p_rating.paragraph_format.space_after = Pt(6)
    r_badge = p_rating.add_run(f"Internal Rating: {rating_label}   |   Decision: {decision}")
    r_badge.bold = True
    r_badge.font.size = Pt(11)
    color = "00700C" if "APPROVED" in decision.upper() else "C00000"
    r_badge.font.color.rgb = RGBColor(*bytes.fromhex(color))

    # View prose
    p_view = doc.add_paragraph()
    p_view.paragraph_format.space_after = Pt(6)
    r_view = p_view.add_run(view_text)
    r_view.font.size = Pt(10)

    # Conditions — plain paragraphs, no header
    if conditions:
        p_cond_intro = doc.add_paragraph()
        p_cond_intro.paragraph_format.space_after = Pt(4)
        r_ci = p_cond_intro.add_run("The approval is subject to the following conditions:")
        r_ci.font.size = Pt(10)
        r_ci.italic = True

        for cond in conditions:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run(cond)
            r.font.size = Pt(10)

    doc.add_paragraph()


# ──────────────────────────── scoring table (end) ────────────────────────────

def add_scoring_table(doc, scoring):
    """
    Change #6: scoring section moved to very end of document.
    """
    heading(doc, "Internal Scoring Summary")

    params = scoring.get("parameters", [])
    total_score = scoring.get("total_score", 0)
    max_score = scoring.get("max_score", 100)
    rating = scoring.get("rating", "")

    table = doc.add_table(rows=1, cols=4)
    set_table_border(table, "AAAAAA", "4")
    for i, h in enumerate(["Parameter", "Max Score", "Score Obtained", "Remarks"]):
        cell = table.cell(0, i)
        set_cell_bg(cell, "1F3864")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for param in params:
        row = table.add_row()
        values = [
            param.get("parameter", ""),
            str(param.get("max_score", "")),
            str(param.get("score", "")),
            param.get("remarks", ""),
        ]
        for i, val in enumerate(values):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.size = Pt(9)

    # Total row
    total_row = table.add_row()
    set_cell_bg(total_row.cells[0], "D9E1F2")
    set_cell_bg(total_row.cells[1], "D9E1F2")
    set_cell_bg(total_row.cells[2], "D9E1F2")
    set_cell_bg(total_row.cells[3], "D9E1F2")
    for i, val in enumerate(["TOTAL", str(max_score), str(total_score), f"Rating: {rating}"]):
        p = total_row.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(val)
        r.bold = True
        r.font.size = Pt(9)

    doc.add_paragraph()


# ──────────────────────────── main generate ────────────────────────────

def generate(data: dict, output_path: str):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(1.5)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    company = data.get("company_name", "Company")
    date_str = data.get("report_date", "")

    # ── Banner ──
    add_banner(doc, company, date_str)

    # ── Change #8: Requirement first, before header table ──
    req = data.get("requirement", "")
    if req:
        p_req = doc.add_paragraph()
        p_req.paragraph_format.space_after = Pt(6)
        r_req_label = p_req.add_run("Requirement: ")
        r_req_label.bold = True
        r_req_label.font.size = Pt(10)
        r_req_label.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        r_req_val = p_req.add_run(req)
        r_req_val.font.size = Pt(10)

    # ── Change #7: 4-field header table ──
    add_header_table(doc, data)

    # ── Initial Screening ──
    if data.get("screening"):
        add_screening(doc, data["screening"])

    # ── Change #1: Brief Profile as prose ──
    if data.get("brief_profile_paragraphs"):
        add_brief_profile(doc, data["brief_profile_paragraphs"])

    # ── Shareholding ──
    if data.get("shareholders"):
        add_shareholding(doc, data["shareholders"])

    # ── Financial Table ──
    if data.get("financials"):
        add_financial_table(doc, data["financials"])

    # ── Change #2: Financial Analysis as prose ──
    if data.get("financial_analysis_sections"):
        add_financial_analysis(doc, data["financial_analysis_sections"])

    # ── Change #3: Strengths ──
    if data.get("strengths"):
        add_sw(doc, "Strengths", data["strengths"])

    # ── Change #3: Weaknesses ──
    if data.get("weaknesses"):
        add_sw(doc, "Weaknesses / Risks", data["weaknesses"])

    # ── Change #4: Latest Updates ──
    if data.get("latest_updates"):
        add_latest_updates(doc, data["latest_updates"])

    # ── Change #5 & #6: Credit View + Conditions (scoring at end) ──
    add_view_and_conditions(
        doc,
        data.get("credit_view", ""),
        data.get("conditions", []),
        data.get("internal_rating", ""),
        data.get("decision", ""),
    )

    # ── Change #6: Scoring table at end ──
    if data.get("scoring"):
        divider(doc)
        add_scoring_table(doc, data["scoring"])

    doc.save(output_path)
    print(f"Report saved → {output_path}")


# ──────────────────────────── entry point ────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python generate_report_v2.py <data.json> <output.docx>")
        sys.exit(1)
    with open(sys.argv[1], encoding="utf-8") as f:
        report_data = json.load(f)
    generate(report_data, sys.argv[2])
